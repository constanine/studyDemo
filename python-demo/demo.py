import xlrd; #xls
import openpyxl; #xlsx
import docx; #docx
from win32com import client #doc


def startXls():
    workbook = xlrd.open_workbook("test-A.xls")
    sheet = workbook.sheet_by_index(0)
    rowSize = sheet.nrowsB
    workbookB = xlrd.open_workbook("test-B.xls")
    sheetB = workbookB.sheet_by_index(0)
    rowSizeB = sheetB.nrows
    for rowIdx in range(rowSize):
        rowData = sheet.row_values(rowIdx)
        rowDataB = sheetB.row_values(rowIdx)
        print(rowDataB[1]+","+rowData[0])
        if (rowDataB[1] == rowData[0]):
            print("相匹配")
        else:
            print("不匹配")

def startXlsx():
    workbook = openpyxl.load_workbook("test-A.xlsx")
    sheet_names = workbook.sheetnames;
    sheet = workbook[sheet_names[0]]
    rowSize = sheet.max_row
    workbookB = openpyxl.load_workbook("test-B.xlsx")
    sheet_namesB = workbookB.sheetnames;
    sheetB = workbookB[sheet_names[0]]
    rowSizeB = sheetB.max_row
    for rowIdx in range(1,rowSize):
        cell = sheet.cell(row = rowIdx,column =1)
        cellB = sheetB.cell(row = rowIdx,column =2)
        print(cell.value+","+cellB.value)
        if (cell.value == cellB.value):
            print("相匹配")
        else:
            print("不匹配")

def startDocx():
    doc = docx.Document("test-A.docx");
    paragraphs = doc.paragraphs;
    for paragraph in paragraphs:
        print(paragraph.text)
    tables = doc.tables;
    #for table in tables:
    #    for row in table.rows:  # 读每行
    #        #row_content = []
    #       for cell in row.cells:  # 读一行中的所有单元格
    #            c = cell.text
    #            print(c)
    #            row_content.append(c)
    #            print(row_content)  # 以列表形式导出每一行数据
    testTxt = tables[0].rows[0].cells[0].text
    print(testTxt)
    if(testTxt.startswith("表1")) :
        print("匹配")
    else :
        print("不匹配")
    if("表2" in testTxt) :
        print("匹配")
    else :
        print("不匹配")

def startDoc():
    word = client.Dispatch('Word.Application')
    doc = word.Documents.Open(FileName="D:/work/GIT_WORK/studyDemo/python-demo/test-A.doc",Encoding='gbk')
    paragraphs = doc.paragraphs;
    for paragraph in paragraphs:
        print(paragraph)
    tables = doc.tables;
    # for table in tables:
    #    for row in table.rows:  # 读每行
    #        #row_content = []
    #       for cell in row.cells:  # 读一行中的所有单元格
    #            c = cell.text
    #            print(c)n
    #            row_content.append(c)
    #            print(row_content)  # 以列表形式导出每一行数据
    table = tables[0]
    testTxt = table.Cell(0, 0).Range.Text
    testTxt=testTxt.replace("\t"," ").replace("\r","").replace("\n"," ")
    print("table-0-0,"+testTxt.encode("utf-8").decode("utf-8"))
    if (testTxt.startswith("表1")):
        print("匹配")
    else:
        print("不匹配")
    if ("表2" in testTxt):
        print("匹配")
    else:
        print("不匹配")
#startXls();
#startDocx();
#startXlsx();
startDoc();