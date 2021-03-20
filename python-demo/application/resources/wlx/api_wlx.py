import xlrd; #xls
import openpyxl; #xlsx
import docx; #docx
from win32com import client #doc
import ExcelTextboxTool;
import DocTextboxTool;
import os;

#查找columnIdxA列中值为sourceVal的右边一格的值
def getRightCellValue(filepath,sheetIdx,columnIdxA,sourceVal):
    return getManualCellValue(filepath,sheetIdx,columnIdxA,columnIdxA+1,sourceVal)

#汇总定义4种格式文件的获取单元格内容（文件地址，第几个sheet/第几个表格，查找第columnIdxA列中值为sourceVal的对应第columnIdxB列中的值）
def getManualCellValue(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal):
    if filepath.endswith(".xlsx"):
        return __getManualCellValueXlsx(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal)
    elif filepath.endswith(".xls"):
        return __getManualCellValueXls(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal)
    elif filepath.endswith(".docx"):
        return __getManualCellValueDocx(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal)
    elif filepath.endswith(".doc"):
        return __getManualCellValueDoc(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal)

#查找第columnIdxA列中值为sourceVal的对应列名为columnTitle中的值
def findCodeVal(filepath,sheetIdx,columnIdxA,columnTitle,sourceVal):
    columnIdxB = findCodeColumnIdx(filepath,sheetIdx,columnTitle);
    return getManualCellValue(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal)

def __findCodeColumnIdxXlsx(filepath, sheetIdx, columnTitle):
    try:
        workbook = xlrd.open_workbook(filepath)
        sheet = workbook.sheet_by_index(sheetIdx)
        cells = sheet['1']
        columnIdx = 0;
        for cell in cells:
            if cell.text == columnTitle:
                return columnIdx
            columnIdx += 1
    finally:
        workbook.release_resources();
        workbook = None;


def __findCodeColumnIdxXls(filepath, sheetIdx, columnTitle):
    try:
        workbook = xlrd.open_workbook(filepath)
        sheet = workbook.sheet_by_index(sheetIdx)
        cells = sheet.row(0);
        columnIdx = 0;
        for cell in cells:
            if cell.text == columnTitle:
                return columnIdx
            columnIdx += 1
    finally:
        workbook.release_resources();
        workbook = None;


def __findCodeColumnIdxDocx(filepath, sheetIdx, columnTitle):
    try:
        doc = docx.Document(filepath);
        tables = doc.tables;
        table = tables[sheetIdx];
        row = table.rows[0];
        columnIdx = 0;
        for cell in row:
            if cell.text == columnTitle:
                return columnIdx
            columnIdx += 1
    finally:
        doc = None;


def __findCodeColumnIdxDoc(filepath, sheetIdx, columnTitle):
    try:
        word = client.Dispatch('Word.Application')
        doc = word.Documents.Open(FileName=filepath, Encoding='gbk')
        tables = doc.tables;
        table = tables[sheetIdx];
        row = table.rows[0];
        columnIdx = 0;
        for cell in row.Cells:
            cellVal = cell.Range.Text;
            cellVal = cellVal[0:len(cellVal) - 2]
            if cellVal == columnTitle:
                return columnIdx
            columnIdx += 1
    finally:
        word = None;


def findCodeColumnIdx(filepath, sheetIdx, columnTitle):
    if filepath.endswith(".xlsx"):
        return __findCodeColumnIdxXlsx(filepath, sheetIdx, columnTitle)
    elif filepath.endswith(".xls"):
        return __findCodeColumnIdxXls(filepath, sheetIdx, columnTitle)
    elif filepath.endswith(".docx"):
        return __findCodeColumnIdxDocx(filepath, sheetIdx, columnTitle)
    elif filepath.endswith(".doc"):
        return __findCodeColumnIdxDoc(filepath, sheetIdx, columnTitle)

def __getManualCellValueXls(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal):
    try:
        workbook = xlrd.open_workbook(filepath)
        sheet = workbook.sheet_by_index(sheetIdx)
        sheet.rich_text_runlist_map
        rowSize = sheet.nrows
        for rowIdx in range(rowSize):
            rowData = sheet.row_values(rowIdx)
            if (str(rowData[columnIdxA]).strip() == sourceVal):
                return rowData[columnIdxB].strip()
    finally:
        workbook.release_resources();
        workbook = None;

def __getManualCellValueXlsx(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal):
    try:
        workbook = openpyxl.load_workbook(filepath, read_only=True)
        sheet_names = workbook.sheetnames;
        sheet = workbook[sheet_names[sheetIdx]]
        rowSize = sheet.max_row+1
        for rowIdx in range(1, rowSize):
            cell = sheet.cell(row=rowIdx, column=columnIdxA)
        if (cell.value == sourceVal):
            return sheet.cell(row=rowIdx, column=columnIdxB).value
    finally:
        workbook.close();
        workbook = None;


def __getManualCellValueDocx(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal):
    try:
        doc = docx.Document(filepath);
        tables = doc.tables;
        table = tables[sheetIdx];
        for row in table.rows:
            if (row.cells[columnIdxA].text == sourceVal):
                return row.cells[columnIdxB].text
    finally:
        doc = None;

def __getManualCellValueDoc(filepath, sheetIdx, columnIdxA,columnIdxB, sourceVal):
    try:
        word = client.Dispatch('Word.Application')
        doc = word.Documents.Open(FileName=filepath,Encoding='gbk')
        tables = doc.tables;
        table = tables[sheetIdx];
        for row in table.rows:
            cellVal = row.Cells[columnIdxA].Range.Text;
            cellVal = cellVal[0:len(cellVal) - 2]
            if (str(cellVal).strip()== str(sourceVal).strip()):
                result = row.Cells[columnIdxB].Range.Text;
                result = result[0:len(result) - 2];
                return str(result).strip();
    finally:
        word = None;

#获得excel里封面页的版本号
def __cover_version(filepath):
    result = ExcelTextboxTool.getXlsTexts(filepath);
    result2 = result[2].partition("/")
    return result2[2]

#核对xls格式的filepath文件内部版本是否一致，和发布单vv报告列表中的版本是否一致
def __check_version_xls(filepath,number,VVreport):
    try:
        chartemp = filepath.rsplit("_", 1)#获得文档中的查询关键字段
        temp2 = chartemp[1]
        temp3 = temp2.rsplit(".", 1)
        chartemp = temp3[0]

        temp1 = __cover_version(filepath)#获取封面上的版本
        
        workbook = xlrd.open_workbook(filepath)
        sheet = workbook.sheet_by_index(number)#修订页为第number+1个sheet的情况
        rowSize = sheet.nrows
        for rowIdx in range(rowSize):
            rowData = sheet.row_values(rowIdx)
            if (str(rowData[4]).strip() == ""):
                rowData2 = sheet.row_values(rowIdx-1)
                temp2 = str(rowData2[4]).strip()#获取修订页上的版本
                break
        if temp1.upper() != temp2.upper():
            print(filepath + "的文件封面页版本和修订记录版本不一致")
        elif temp1.upper() != getManualCellValue(VVreport,0,0,2,chartemp).upper():
            print(filepath + "的文件版本和VV报告中的版本不一致")

        #补充和VV报告版本一致性的代码

    finally:
        workbook.close();
        workbook = None;

#（文件地址，第几个sheet/第几个表格，查找第columnIdxA列中值为以sourceVal开头的对应第columnIdxB列中的值）
def getManualCellValuestart(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal):
    if filepath.endswith(".xls"):
        return __getManualCellValueXlsstart(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal)
    elif filepath.endswith(".doc"):
        return __getManualCellValueDocstart(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal)

def __getManualCellValueXlsstart(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal):
    try:
        workbook = xlrd.open_workbook(filepath)
        sheet = workbook.sheet_by_index(sheetIdx)
        sheet.rich_text_runlist_map
        rowSize = sheet.nrows
        for rowIdx in range(rowSize):
            rowData = sheet.row_values(rowIdx)
            if (str(rowData[columnIdxA]).strip().startswith(sourceVal)):
                return rowData[columnIdxB].strip()
    finally:
        workbook.release_resources();
        workbook = None;

def __getManualCellValueDocstart(filepath, sheetIdx, columnIdxA,columnIdxB, sourceVal):
    try:
        word = client.Dispatch('Word.Application')
        doc = word.Documents.Open(FileName=filepath,Encoding='gbk')
        tables = doc.tables;
        table = tables[sheetIdx];
        for row in table.rows:
            cellVal = row.Cells[columnIdxA].Range.Text;
            cellVal = cellVal[0:len(cellVal) - 2]
            if (str(cellVal).strip().startswith(sourceVal)):
                result = row.Cells[columnIdxB].Range.Text;
                result = result[0:len(result) - 2];
                return str(result).strip();
    finally:
        word = None;

def _checkfileexist(title):
    pathname = input(title)
    while 1:
        if os.path.exists(pathname) == True:
            return pathname
        else:
            pathname = input(title)