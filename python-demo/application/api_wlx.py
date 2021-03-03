import xlrd; #xls
import openpyxl; #xlsx
import docx; #docx
from win32com import client #doc




def getRightCellValue(filepath,sheetIdx,columnIdxA,sourceVal):
    return getManualCellValue(filepath,sheetIdx,columnIdxA,columnIdxA+1,sourceVal)

def getManualCellValue(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal):
    if filepath.endswith(".xlsx"):
        return __getManualCellValueXlsx(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal)
    elif filepath.endswith(".xls"):
        return __getManualCellValueXls(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal)
    elif filepath.endswith(".docx"):
        return __getManualCellValueDocx(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal)
    elif filepath.endswith(".doc"):
        return __getManualCellValueDoc(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal)

def findCodeVal(filepath,sheetIdx,columnIdxA,columnTitle,sourceVal):
    columnIdxB = findCodeColumnIdx(filepath,sheetIdx,columnTitle);
    return getManualCellValue(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal)

def __findCodeColumnIdxXlsx(filepath, sheetIdx, columnTitle):
    try:
        workbook = xlrd.open_workbook(filepath)
        sheet = workbook.sheet_by_index(sheetIdx)
        cells = sheet['1']
        columnIdx = 0;
        for cell in cells:
            if cell.text == columnTitle:
                return columnIdx
            columnIdx += 1
    finally:
        workbook.release_resources();
        workbook = None;


def __findCodeColumnIdxXls(filepath, sheetIdx, columnTitle):
    try:
        workbook = xlrd.open_workbook(filepath)
        sheet = workbook.sheet_by_index(sheetIdx)
        cells = sheet.row(0);
        columnIdx = 0;
        for cell in cells:
            if cell.text == columnTitle:
                return columnIdx
            columnIdx += 1
    finally:
        workbook.release_resources();
        workbook = None;


def __findCodeColumnIdxDocx(filepath, sheetIdx, columnTitle):
    try:
        doc = docx.Document(filepath);
        tables = doc.tables;
        table = tables[sheetIdx];
        row = table.rows[0];
        columnIdx = 0;
        for cell in row:
            if cell.text == columnTitle:
                return columnIdx
            columnIdx += 1
    finally:
        doc = None;


def __findCodeColumnIdxDoc(filepath, sheetIdx, columnTitle):
    try:
        word = client.Dispatch('Word.Application')
        doc = word.Documents.Open(FileName=filepath, Encoding='gbk')
        tables = doc.tables;
        table = tables[sheetIdx];
        row = table.rows[0];
        columnIdx = 0;
        for cell in row.Cells:
            cellVal = cell.Range.Text;
            cellVal = cellVal[0:len(cellVal) - 2]
            if cellVal == columnTitle:
                return columnIdx
            columnIdx += 1
    finally:
        word = None;


def findCodeColumnIdx(filepath, sheetIdx, columnTitle):
    if filepath.endswith(".xlsx"):
        return __findCodeColumnIdxXlsx(filepath, sheetIdx, columnTitle)
    elif filepath.endswith(".xls"):
        return __findCodeColumnIdxXls(filepath, sheetIdx, columnTitle)
    elif filepath.endswith(".docx"):
        return __findCodeColumnIdxDocx(filepath, sheetIdx, columnTitle)
    elif filepath.endswith(".doc"):
        return __findCodeColumnIdxDoc(filepath, sheetIdx, columnTitle)

def __getManualCellValueXls(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal):
    try:
        workbook = xlrd.open_workbook(filepath)
        sheet = workbook.sheet_by_index(sheetIdx)
        sheet.rich_text_runlist_map
        rowSize = sheet.nrows
        for rowIdx in range(rowSize):
            rowData = sheet.row_values(rowIdx)
            if (rowData[columnIdxA] == sourceVal):
                return rowData[columnIdxB]
    finally:
        workbook.release_resources();
        workbook = None;

def __getManualCellValueXlsx(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal):
    try:
        workbook = openpyxl.load_workbook(filepath, read_only=True)
        sheet_names = workbook.sheetnames;
        sheet = workbook[sheet_names[sheetIdx]]
        rowSize = sheet.max_row+1
        for rowIdx in range(1, rowSize):
            cell = sheet.cell(row=rowIdx, column=columnIdxA)
        if (cell.value == sourceVal):
            return sheet.cell(row=rowIdx, column=columnIdxB).value
    finally:
        workbook.close();
        workbook = None;


def __getManualCellValueDocx(filepath,sheetIdx,columnIdxA,columnIdxB,sourceVal):
    try:
        doc = docx.Document(filepath);
        tables = doc.tables;
        table = tables[sheetIdx];
        for row in table.rows:
            if (row.cells[columnIdxA].text == sourceVal):
                return row.cells[columnIdxB].text
    finally:
        doc = None;

def __getManualCellValueDoc(filepath, sheetIdx, columnIdxA,columnIdxB, sourceVal):
    try:
        word = client.Dispatch('Word.Application')
        doc = word.Documents.Open(FileName=filepath,Encoding='gbk')
        tables = doc.tables;
        table = tables[sheetIdx];
        for row in table.rows:
            cellVal = row.Cells[columnIdxA].Range.Text;
            cellVal = cellVal[0:len(cellVal) - 2]
            if (cellVal == sourceVal):
                result = row.Cells[columnIdxB].Range.Text;
                result = result[0:len(result) - 2];
                return result;
    finally:
        word = None;